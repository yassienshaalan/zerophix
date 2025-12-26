import io
from typing import Dict, Any, Optional, List
from abc import ABC, abstractmethod
import logging

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

class DocumentProcessor(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    def extract_text(self, content: bytes) -> str:
        """Extract text from document content"""
        pass
    
    @abstractmethod
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Redact text and reconstruct document with formatting preserved"""
        pass
    
    @abstractmethod
    def get_supported_types(self) -> List[str]:
        """Get list of supported MIME types"""
        pass

class PlainTextProcessor(DocumentProcessor):
    """Processor for plain text files"""
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from plain text content"""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file")
    
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Apply redactions to plain text"""
        text = self.extract_text(content)
        
        # Apply redactions
        for original, redacted in redaction_map.items():
            text = text.replace(original, redacted)
        
        return text.encode('utf-8')
    
    def get_supported_types(self) -> List[str]:
        return ["text/plain"]

class PDFProcessor(DocumentProcessor):
    """Processor for PDF files"""
    
    def __init__(self):
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
                except Exception as e:
                    logging.warning(f"Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            return text
        except Exception as e:
            raise ValueError(f"Could not extract text from PDF: {e}")
    
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Redact PDF content (simplified - extracts text and returns as PDF)"""
        # Note: Full PDF redaction with formatting preservation requires more complex libraries
        # like reportlab or pypdf2 with additional processing
        
        text = self.extract_text(content)
        
        # Apply redactions
        for original, redacted in redaction_map.items():
            text = text.replace(original, redacted)
        
        # For now, return as text (in production, would reconstruct PDF)
        return text.encode('utf-8')
    
    def get_supported_types(self) -> List[str]:
        return ["application/pdf"]

class DOCXProcessor(DocumentProcessor):
    """Processor for Microsoft Word DOCX files"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX processing requires python-docx. Install with: pip install python-docx")
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
            
            return text
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {e}")
    
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Redact DOCX content with formatting preservation"""
        try:
            doc = Document(io.BytesIO(content))
            
            # Redact paragraphs
            for paragraph in doc.paragraphs:
                for original, redacted in redaction_map.items():
                    if original in paragraph.text:
                        # Simple replacement (in production, would preserve formatting)
                        paragraph.text = paragraph.text.replace(original, redacted)
            
            # Redact tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for original, redacted in redaction_map.items():
                            if original in cell.text:
                                cell.text = cell.text.replace(original, redacted)
            
            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            raise ValueError(f"Could not redact DOCX: {e}")
    
    def get_supported_types(self) -> List[str]:
        return ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

class ExcelProcessor(DocumentProcessor):
    """Processor for Excel files"""
    
    def __init__(self):
        if not EXCEL_AVAILABLE:
            raise RuntimeError("Excel processing requires openpyxl and pandas. Install with: pip install openpyxl pandas")
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from Excel content"""
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(io.BytesIO(content))
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                # Convert DataFrame to text representation
                text += df.to_string(index=False, na_rep='')
                text += "\n"
            
            return text
        except Exception as e:
            raise ValueError(f"Could not extract text from Excel: {e}")
    
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Redact Excel content"""
        try:
            excel_file = pd.ExcelFile(io.BytesIO(content))
            
            # Create new workbook
            with pd.ExcelWriter(io.BytesIO(), engine='openpyxl') as writer:
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    # Apply redactions to all string columns
                    for col in df.select_dtypes(include=['object']).columns:
                        for original, redacted in redaction_map.items():
                            df[col] = df[col].astype(str).str.replace(original, redacted, regex=False)
                    
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                return writer.book
                
        except Exception as e:
            raise ValueError(f"Could not redact Excel: {e}")
    
    def get_supported_types(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
            "application/vnd.ms-excel"  # .xls
        ]

class CSVProcessor(DocumentProcessor):
    """Processor for CSV files"""
    
    def extract_text(self, content: bytes) -> str:
        """Extract text from CSV content"""
        try:
            text_content = content.decode('utf-8')
            
            # Parse CSV and convert to readable format
            csv_reader = csv.reader(io.StringIO(text_content))
            formatted_text = ""
            
            for row_num, row in enumerate(csv_reader):
                if row_num == 0:  # Header row
                    formatted_text += "Headers: " + " | ".join(row) + "\n"
                    formatted_text += "-" * 50 + "\n"
                else:
                    formatted_text += " | ".join(row) + "\n"
            
            return formatted_text
        except Exception as e:
            raise ValueError(f"Could not extract text from CSV: {e}")
    
    def redact_and_reconstruct(self, content: bytes, redaction_map: Dict[str, str]) -> bytes:
        """Redact CSV content"""
        try:
            text_content = content.decode('utf-8')
            
            # Apply redactions to the entire content
            for original, redacted in redaction_map.items():
                text_content = text_content.replace(original, redacted)
            
            return text_content.encode('utf-8')
        except Exception as e:
            raise ValueError(f"Could not redact CSV: {e}")
    
    def get_supported_types(self) -> List[str]:
        return ["text/csv", "application/csv"]

class DocumentProcessorFactory:
    """Factory for creating document processors"""
    
    def __init__(self):
        self.processors = {}
        self._register_processors()
    
    def _register_processors(self):
        """Register available document processors"""
        # Always available
        self.processors["text/plain"] = PlainTextProcessor
        self.processors["text/csv"] = CSVProcessor
        self.processors["application/csv"] = CSVProcessor
        
        # Optional processors based on available libraries
        if PDF_AVAILABLE:
            self.processors["application/pdf"] = PDFProcessor
        
        if DOCX_AVAILABLE:
            self.processors["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] = DOCXProcessor
        
        if EXCEL_AVAILABLE:
            self.processors["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"] = ExcelProcessor
            self.processors["application/vnd.ms-excel"] = ExcelProcessor
    
    def get_processor(self, content_type: str) -> Optional[DocumentProcessor]:
        """Get processor for given content type"""
        processor_class = self.processors.get(content_type)
        if processor_class:
            try:
                return processor_class()
            except RuntimeError as e:
                logging.warning(f"Could not create processor for {content_type}: {e}")
                return None
        return None
    
    def get_supported_types(self) -> List[str]:
        """Get list of all supported content types"""
        return list(self.processors.keys())
    
    def is_supported(self, content_type: str) -> bool:
        """Check if content type is supported"""
        return content_type in self.processors

class DocumentRedactionService:
    """High-level service for document redaction"""
    
    def __init__(self, redaction_pipeline):
        """
        Initialize document redaction service
        
        Args:
            redaction_pipeline: RedactionPipeline instance
        """
        self.redaction_pipeline = redaction_pipeline
        self.processor_factory = DocumentProcessorFactory()
    
    async def redact_document(self, content: bytes, content_type: str, 
                            preserve_formatting: bool = True) -> Dict[str, Any]:
        """
        Redact a document while optionally preserving formatting
        
        Args:
            content: Document content as bytes
            content_type: MIME type of the document
            preserve_formatting: Whether to preserve original formatting
        
        Returns:
            Dictionary with redaction results
        """
        # Get appropriate processor
        processor = self.processor_factory.get_processor(content_type)
        if not processor:
            raise ValueError(f"Unsupported document type: {content_type}")
        
        try:
            # Extract text from document
            extracted_text = processor.extract_text(content)
            
            # Perform redaction on extracted text
            redaction_result = self.redaction_pipeline.redact(extracted_text)
            
            # Create redaction mapping
            redaction_map = {}
            original_text = extracted_text
            redacted_text = redaction_result["text"]
            
            # Build mapping from spans
            for span in redaction_result.get("spans", []):
                original_entity = original_text[span["start"]:span["end"]]
                # Find corresponding redacted text (simplified)
                redaction_map[original_entity] = f"<{span['label']}>"
            
            # Reconstruct document if formatting preservation is requested
            if preserve_formatting:
                try:
                    redacted_content = processor.redact_and_reconstruct(content, redaction_map)
                    output_format = "preserved"
                except Exception as e:
                    logging.warning(f"Could not preserve formatting: {e}. Returning plain text.")
                    redacted_content = redacted_text.encode('utf-8')
                    output_format = "plain_text"
            else:
                redacted_content = redacted_text.encode('utf-8')
                output_format = "plain_text"
            
            return {
                "success": True,
                "original_size": len(content),
                "redacted_size": len(redacted_content),
                "redacted_content": redacted_content,
                "output_format": output_format,
                "entities_found": len(redaction_result.get("spans", [])),
                "redaction_stats": redaction_result.get("stats", {}),
                "extracted_text_length": len(extracted_text),
                "redacted_text_length": len(redacted_text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "content_type": content_type
            }
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get information about supported formats"""
        supported_types = self.processor_factory.get_supported_types()
        
        format_info = {}
        for content_type in supported_types:
            processor = self.processor_factory.get_processor(content_type)
            if processor:
                format_info[content_type] = {
                    "extract_text": True,
                    "preserve_formatting": hasattr(processor, 'redact_and_reconstruct'),
                    "processor_class": processor.__class__.__name__
                }
        
        return format_info